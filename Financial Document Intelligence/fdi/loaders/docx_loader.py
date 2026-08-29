from docx import Document as DocxFile

from fdi.schema import Document, Table


def load_docx(path: str, category: str) -> Document:
    """Extract paragraph text and tables from a Word document into a normalized Document."""
    docx = DocxFile(path)
    text = "\n".join(p.text for p in docx.paragraphs)

    tables: list[Table] = []
    for table in docx.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])

    return Document(
        source_path=path,
        category=category,
        file_type="docx",
        text=text,
        tables=tables,
    )
